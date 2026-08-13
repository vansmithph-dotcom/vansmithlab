# -*- coding: utf-8 -*-
"""Bring new Analysis DOCX files up to DOCX_SCHEMA v1.3.

Fixes per file:
  1. HERO line 2 must start with "Subtitle: ".
  2. Every [SECTION id="sec-N"] heading paragraph gets a Word bookmark
     named after the section id.
  3. Each [TOC] entry becomes a Word hyperlink to the matching bookmark.

The converter reads paragraph text, so these changes do not alter the body
markdown; they only satisfy scripts/validate-structure.py.
"""
import sys, re, os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FILES = [
    "VSL_Visionaire_The_Magazine_That_Refused_to_Be_a_Magazine_RU.docx",
    "VSL_Visionaire_The_Magazine_That_Refused_to_Be_a_Magazine_EN.docx",
    "VSL_Why_Dazed_and_iD_Matter_Beyond_Youth_Magazines_RU.docx",
    "VSL_Why_Dazed_and_iD_Matter_Beyond_Youth_Magazines_EN.docx",
    "VSL_Magazine_Is_No_Longer_Paper_RU.docx",
    "VSL_Magazine_Is_No_Longer_Paper_EN.docx",
    "VSL_Why_Fashion_House_Now_Resembles_TV_Network_RU.docx",
    "VSL_Why_Fashion_House_Now_Resembles_TV_Network_EN.docx",
    "VSL_Grace_Coddington_Fashion_Editor_Narrative_Image_Making_RU.docx",
    "VSL_Grace_Coddington_Fashion_Editor_Narrative_Image_Making_EN.docx",
]

def ptext(p):
    return "".join(n.text or "" for n in p._p.iter(qn("w:t")))

def set_ptext(p, text):
    # remove existing runs and set one run
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.add_run(text)

def add_bookmark(p, name):
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(abs(hash(name)) % 100000))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), bm_start.get(qn("w:id")))
    p._p.insert(0, bm_start)
    p._p.append(bm_end)

def add_hyperlink(p, anchor, text):
    # create a run wrapped in a hyperlink
    h = OxmlElement("w:hyperlink")
    h.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    h.append(r)
    # clear existing runs then append hyperlink
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p._p.append(h)

def fix(path):
    doc = Document(path)
    paras = doc.paragraphs
    texts = [ptext(p).strip() for p in paras]

    # 1. HERO subtitle
    for i, t in enumerate(texts):
        if t == "[HERO]":
            if i + 2 < len(paras):
                sub = ptext(paras[i + 2]).strip()
                if sub and not sub.lower().startswith("subtitle:"):
                    set_ptext(paras[i + 2], "Subtitle: " + sub)
            break

    # 2. bookmarks on section headings
    sec_name = None
    for p, t in zip(paras, texts):
        m = re.fullmatch(r'\[SECTION id="([^"]+)"\]', t)
        if m:
            sec_name = m.group(1)
            continue
        if t == "[/SECTION]":
            sec_name = None
            continue
        if sec_name:
            add_bookmark(p, sec_name)
            sec_name = None

    # 3. TOC hyperlinks
    in_toc = False
    for p, t in zip(paras, texts):
        if t == "[TOC]":
            in_toc = True
            continue
        if t == "[/TOC]":
            in_toc = False
            continue
        if in_toc:
            m = re.match(r"^(\d+)\.\s+(.*)$", t)
            if m:
                add_hyperlink(p, "sec-" + m.group(1), t)

    doc.save(path)
    print("fixed", os.path.basename(path))

if __name__ == "__main__":
    d = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\VAN\Downloads"
    for fn in FILES:
        p = os.path.join(d, fn)
        if not os.path.exists(p):
            print("skip (not present)", fn)
            continue
        fix(p)
    print("done")
